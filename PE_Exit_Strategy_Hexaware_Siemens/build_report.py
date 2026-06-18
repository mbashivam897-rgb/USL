#!/usr/bin/env python3
"""Generate the PE Exits & M&A case-study report as a .docx file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x6E, 0xA8)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ----- base styles -----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY if level == 1 else ACCENT
    run.font.size = Pt(16 if level == 1 else 13)
    if level == 1:
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "1F3A5F")
        pbdr.append(bottom)
        pPr.append(pbdr)
    return p


def subhead(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = GREY
    return p


def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(3)
    return p


def make_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "1F3A5F")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(10)


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(3):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Private Equity Exit Strategies and\nM&A Value Creation")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("A Case Study of Hexaware Technologies, Siemens India,\nand the HDFC\u2013HDFC Bank Merger")
r.font.size = Pt(15)
r.font.color.rgb = ACCENT
r.italic = True

for _ in range(4):
    doc.add_paragraph()

for label in [
    "A Case-Study-Based Project in Finance",
    "Submitted in partial fulfilment of the requirements of the",
    "Master of Business Administration (MBA) \u2013 Finance",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.font.size = Pt(12)
    r.font.color.rgb = GREY

for _ in range(6):
    doc.add_paragraph()

for label in ["Submitted by: ____________________",
              "Roll No.: ____________________",
              "Guided by: ____________________",
              "Academic Year: 2025\u201326"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
heading("Table of Contents", 1)
toc_items = [
    ("Executive Summary", ""),
    ("Section 1: Introduction to Private Equity and M&A", ""),
    ("   1.1 Overview of Private Equity", ""),
    ("   1.2 The PE Investment Lifecycle", ""),
    ("   1.3 Exit Strategies in Private Equity", ""),
    ("   1.4 The Role of M&A in Value Creation", ""),
    ("Section 2: Case Study \u2013 Hexaware Technologies (Carlyle Group)", ""),
    ("Section 3: Case Study \u2013 Siemens India", ""),
    ("Section 4: M&A Case Study \u2013 HDFC Ltd & HDFC Bank Merger", ""),
    ("Section 5: Comparative Private Equity Exit Analysis", ""),
    ("Section 6: Key Findings and Recommendations", ""),
    ("References and Data Sources", ""),
]
for item, _ in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(item)
    run.font.size = Pt(11)
    if not item.startswith("   "):
        run.bold = True
doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
heading("Executive Summary", 1)
body(
    "This project examines how private equity (PE) investors create value during their holding "
    "period and realise returns through well-timed exits, and how mergers and acquisitions (M&A) "
    "serve as instruments of strategic value creation. Rather than building a discounted-cash-flow "
    "valuation, the study adopts a case-based approach, analysing three landmark Indian "
    "transactions that together illustrate the full spectrum of exit and consolidation strategies."
)
body(
    "The first case studies the Carlyle Group's acquisition of Hexaware Technologies in 2021 \u2014 "
    "at roughly USD 3 billion, India's largest private-equity deal at the time \u2014 and Carlyle's "
    "subsequent partial exit through a Rs 8,750-crore initial public offering (IPO) in February 2025, "
    "the largest IT-services IPO in Indian history. The second case is a forward-looking, analytical "
    "exercise on Siemens India, evaluating which exit and M&A routes would be most attractive were a "
    "controlling stake ever offered to financial or strategic buyers, set against the backdrop of the "
    "company's 2024\u201325 demerger of its energy business. The third case dissects the HDFC Ltd\u2013HDFC "
    "Bank merger \u2014 a roughly USD 40 billion combination effective 1 July 2023 that created the "
    "world's fourth-largest bank by market capitalisation \u2014 as a textbook study in synergy-driven M&A."
)
body(
    "Across the three cases the report draws out a consistent set of lessons: value creation depends "
    "on a clear investment thesis, operational improvement, disciplined capital structure and, above "
    "all, exit timing aligned with favourable market windows. The study closes with a comparative "
    "framework of the five principal PE exit routes and a recommended exit strategy for Siemens India."
)

# ============================================================
# SECTION 1
# ============================================================
heading("Section 1: Introduction to Private Equity and M&A", 1)

subhead("1.1 Overview of Private Equity")
body(
    "Private equity refers to capital invested in companies that are not listed on a public stock "
    "exchange, or in public companies with the intention of taking them private. PE funds raise "
    "capital from institutional investors and high-net-worth individuals (limited partners), pool it "
    "into a fund managed by a general partner, and deploy it into target companies with the objective "
    "of improving their performance and selling the stake at a profit, typically over a three- to "
    "seven-year horizon."
)
body(
    "Returns to PE investors are commonly measured using two metrics. The Multiple on Invested "
    "Capital (MOIC) expresses total value returned as a multiple of capital invested, while the "
    "Internal Rate of Return (IRR) captures the annualised, time-weighted return and therefore "
    "rewards a quicker exit."
)

subhead("1.2 The PE Investment Lifecycle")
numbered("Sourcing and due diligence \u2014 identifying targets and validating the investment thesis.")
numbered("Acquisition and structuring \u2014 negotiating price, deal structure and financing.")
numbered("Value creation (the holding period) \u2014 operational improvement, governance upgrades, "
         "bolt-on acquisitions and margin expansion.")
numbered("Exit \u2014 monetising the investment through a sale, IPO or recapitalisation.")

subhead("1.3 Exit Strategies in Private Equity")
body("PE funds typically realise returns through one of the following exit routes:")
bullet("Initial Public Offering (IPO): listing the company on a stock exchange and selling shares "
       "to public investors, often in stages.")
bullet("Strategic sale (trade sale): selling to an industry buyer that can extract synergies and "
       "is often willing to pay a control premium.")
bullet("Secondary buyout / sponsor-to-sponsor sale: selling to another PE fund.")
bullet("Partial exit: selling a portion of the holding (for example through an IPO offer-for-sale) "
       "while retaining residual upside.")
bullet("Recapitalisation: returning capital to investors through additional leverage while "
       "retaining ownership.")

subhead("1.4 The Role of M&A in Value Creation")
body(
    "Mergers and acquisitions create value when the combined entity is worth more than the sum of "
    "its parts. The principal sources of value are revenue synergies (cross-selling, wider "
    "distribution), cost synergies (scale economies, lower funding costs, removal of duplicate "
    "functions), financial synergies (a stronger balance sheet and improved credit profile) and "
    "strategic repositioning. For PE investors, M&A is both an entry tool (buy-and-build platforms) "
    "and an exit tool (a trade sale to a strategic acquirer)."
)

# ============================================================
# SECTION 2 - HEXAWARE
# ============================================================
heading("Section 2: Case Study \u2013 Hexaware Technologies (Carlyle Group)", 1)

subhead("2.1 Company Overview")
body(
    "Hexaware Technologies is a global provider of IT and business-process services, focused on "
    "digital transformation, cloud, data and AI-led automation. Headquartered in Mumbai with a "
    "significant delivery presence across India and a strong client base in North America and "
    "Europe, Hexaware operates in the mid-tier of the Indian IT services industry."
)

subhead("2.2 Industry Overview")
body(
    "The Indian IT-services sector is among the country's largest export industries, competing "
    "globally on the strength of skilled talent, cost arbitrage and increasingly on digital "
    "engineering capability. Mid-tier players such as Hexaware compete with both the large-cap "
    "leaders and a long tail of niche specialists, differentiating through domain focus, faster "
    "decision-making and aggressive adoption of automation and AI."
)

subhead("2.3 Ownership Timeline and the Carlyle Investment Thesis")
make_table(
    ["Date", "Event"],
    [
        ["November 2020", "Baring Private Equity Asia (BPEA) delists Hexaware from Indian exchanges."],
        ["October\u2013November 2021", "Carlyle Group, via CA Magnum Holdings, acquires a controlling "
         "stake (~95%) from BPEA for ~USD 3 billion (avg ~Rs 385/share) \u2014 India's largest PE deal at the time."],
        ["2021\u20132024", "Holding period: value-creation programme across growth, margins and AI capability."],
        ["February 2025", "Partial exit via IPO: Rs 8,750-crore offer-for-sale; relisting on BSE and NSE."],
    ],
    widths=[1.6, 4.4],
)
caption("Sources: Reuters, Moneycontrol, CNBC-TV18, NDTV Profit, Carlyle press release.")
body(
    "Carlyle's thesis rested on acquiring a proven mid-tier IT-services platform during a period of "
    "accelerating enterprise digital spending, and scaling it through investment in sales, new "
    "service lines (notably AI and automation) and selective acquisitions \u2014 with a public listing "
    "as the natural exit route once growth and profitability had been demonstrated."
)

subhead("2.4 Acquisition, Delisting and Deal Structure")
body(
    "The investment was made through the holding vehicle CA Magnum Holdings. Because Hexaware was "
    "already private (BPEA having taken it off the exchanges in 2020), Carlyle acquired a private "
    "company outright rather than launching a public takeover. This gave Carlyle full operational "
    "control and a clean platform to execute its value-creation plan without the disclosure and "
    "minority-shareholder constraints of a listed entity."
)

subhead("2.5 Value Creation During Carlyle's Ownership")
bullet("Revenue growth: consolidated revenue rose from roughly Rs 7,178 crore in CY2021 to about "
       "Rs 13,430 crore in CY2025, a compound annual growth rate of roughly 17%.")
bullet("Profitability: EBITDA and earnings expanded alongside revenue; CY2024 revenue reached "
       "USD 1,429 million (up 13.7% year-on-year) and CY2025 revenue USD 1,537.4 million (up 7.6%), "
       "with profit after tax of USD 157 million.")
bullet("Capability building: heavy investment in AI-led delivery, automation platforms and "
       "go-to-market expansion in Europe and the banking vertical.")
bullet("Governance: installation of a new board and professional leadership aligned with public-"
       "market readiness.")

subhead("2.6 Financial Performance Trajectory")
make_table(
    ["Calendar Year", "Revenue (Rs crore, approx.)", "Commentary"],
    [
        ["CY2021", "7,178", "Year of Carlyle acquisition"],
        ["CY2022", "9,200", "Strong post-acquisition growth"],
        ["CY2023", "10,380", "Continued double-digit growth"],
        ["CY2024", "11,974", "Revenue USD 1,429m, +13.7% YoY"],
        ["CY2025", "13,430", "Revenue USD 1,537.4m, +7.6% YoY; PAT USD 157m"],
    ],
    widths=[1.6, 2.4, 2.0],
)
caption("Revenue figures in Rs crore are approximate, converted/rounded from reported INR-million "
        "data (Stockopedia, company results, PRNewswire).")

subhead("2.7 The Exit: IPO and Re-listing Strategy")
body(
    "Carlyle chose a partial exit via IPO as its primary monetisation route. In February 2025, "
    "Hexaware launched a Rs 8,750-crore (~USD 1 billion) IPO structured entirely as an offer-for-sale "
    "by the Carlyle affiliate CA Magnum Holdings, meaning the proceeds went to the selling shareholder "
    "rather than the company. The issue was priced in a band of Rs 674\u2013708 per share, valuing the "
    "company at roughly Rs 43,247 crore at the upper end. The IPO opened on 12 February and the shares "
    "listed on 19 February 2025 at about a 5% premium, giving a market capitalisation of around "
    "Rs 44,422 crore."
)
body(
    "The transaction was the largest IT-services IPO in Indian history and, according to Carlyle, the "
    "largest technology-services IPO globally in over a decade. Crucially, because it was an offer-"
    "for-sale and Carlyle retained a substantial residual holding, the listing represented a partial "
    "exit \u2014 monetising part of the gain while preserving exposure to future upside and the option of "
    "further sell-downs."
)

subhead("2.8 Estimated Returns to Carlyle")
body(
    "Carlyle's entry was at approximately USD 3 billion for its controlling stake in 2021. By the "
    "February 2025 listing the company was valued at roughly Rs 43,000\u201344,000 crore (broadly "
    "USD 5\u20135.4 billion). The IPO alone returned about USD 1 billion in cash to Carlyle while it "
    "continued to hold the majority of the business. The combination of a higher enterprise value and "
    "a large cash realisation points to a clearly profitable outcome; precise MOIC and IRR depend on "
    "the exact retained stake and any subsequent sell-downs, which are not fully public."
)
make_table(
    ["Metric", "Entry (2021)", "At IPO (Feb 2025)"],
    [
        ["Implied equity value", "~USD 3.0 bn (for controlling stake)", "~USD 5.0\u20135.4 bn (whole company)"],
        ["Cash realised by Carlyle", "\u2014", "~USD 1.0 bn (offer-for-sale)"],
        ["Holding status", "~95% control", "Majority retained (partial exit)"],
    ],
    widths=[1.9, 2.0, 2.1],
)
caption("Indicative figures compiled from public reporting; not an official return disclosure.")

subhead("2.9 Key Lessons from the Hexaware Case")
bullet("A clear digital-growth thesis plus operational scaling can roughly double revenue over a "
       "four-year hold.")
bullet("The IPO offer-for-sale structure allows a sponsor to take cash off the table while retaining "
       "upside \u2014 a textbook partial exit.")
bullet("Exit timing into a strong primary-market window materially improves realised value.")

# ============================================================
# SECTION 3 - SIEMENS
# ============================================================
heading("Section 3: Case Study \u2013 Siemens India", 1)
body(
    "Note on framing: Siemens Limited (Siemens India) is majority-owned by its strategic parent, "
    "Siemens AG, not by a private-equity fund. This section is therefore a forward-looking, "
    "analytical exercise \u2014 it evaluates the exit and M&A routes that would be most relevant were a "
    "controlling or large minority stake ever to change hands, and what makes the company attractive "
    "to financial and strategic buyers."
)

subhead("3.1 Company Overview")
body(
    "Siemens Limited is one of India's leading technology and engineering companies, operating in "
    "industrial automation, smart infrastructure, mobility (rail) and, until its recent demerger, "
    "energy. Following the board's 2024 decision to demerge the energy business into a separately "
    "listed entity (Siemens Energy India Limited, on a 1:1 share-entitlement basis), Siemens Limited "
    "now focuses on three core segments: Digital Industries, Smart Infrastructure and Mobility."
)

subhead("3.2 Capital-Goods Industry Analysis and Growth Drivers")
bullet("Government capital expenditure on infrastructure, railways and metro projects.")
bullet("Manufacturing localisation and the 'Make in India' / production-linked-incentive push.")
bullet("Grid modernisation, electrification and the energy transition.")
bullet("Factory automation and digitalisation of industry.")
body(
    "Recent results illustrate both the opportunity and its cyclicality: order inflows have grown "
    "strongly (driven by Digital Industries and Smart Infrastructure), even in quarters where revenue "
    "was affected by a slowdown in short-cycle private-sector capex and demand normalisation."
)

subhead("3.3 Strategic Positioning and Attractiveness to Investors")
body(
    "Siemens India combines a premium brand, a large installed base, a long order backlog and "
    "exposure to multiple structural growth themes. These qualities make it attractive in principle "
    "to two buyer groups:"
)
subhead("Potential strategic buyers")
bullet("Global industrial and automation majors seeking scaled India exposure.")
bullet("Domestic capital-goods and infrastructure conglomerates pursuing consolidation.")
subhead("Potential financial buyers")
bullet("Large global PE funds and infrastructure funds with the capacity to write multi-billion-"
       "dollar cheques and a long investment horizon.")
bullet("Sovereign wealth and pension funds seeking stable, infrastructure-linked cash flows.")

subhead("3.4 Exit / Transaction Routes Available")
make_table(
    ["Route", "Description", "Suitability for Siemens India"],
    [
        ["Demerger + listing", "Separate and list a business unit (as done with Energy)",
         "High \u2014 already proven; unlocks segment value"],
        ["Strategic sale", "Sale of a stake to an industry buyer",
         "Moderate \u2014 limited by parent's strategic intent"],
        ["Secondary stake sale", "Parent dilutes via block deals / OFS",
         "Moderate-High \u2014 improves free float and liquidity"],
        ["Financial-buyer carve-out", "Sell a non-core unit to a PE fund",
         "Moderate \u2014 viable for non-core segments"],
    ],
    widths=[1.5, 2.3, 2.2],
)

subhead("3.5 Most Suitable Strategy, Risks and Opportunities")
body(
    "The energy-business demerger demonstrates that value-unlocking through separation and listing is "
    "the most natural and shareholder-friendly route for a company of Siemens India's profile. A "
    "full PE-style buyout is unlikely given the strategic parent's control, but selective carve-outs "
    "of non-core units to financial buyers, or further demergers, represent the most realistic "
    "value-creation transactions. Key risks include capex cyclicality, commodity-cost and forex "
    "pressure, and execution risk in large projects; the principal opportunity is sustained "
    "infrastructure and electrification spending."
)

# ============================================================
# SECTION 4 - HDFC MERGER
# ============================================================
heading("Section 4: M&A Case Study \u2013 HDFC Ltd & HDFC Bank Merger", 1)

subhead("4.1 Transaction Background")
body(
    "On 4 April 2022, HDFC Bank \u2014 India's largest private-sector bank \u2014 announced that it would "
    "merge with its largest shareholder and parent, Housing Development Finance Corporation (HDFC "
    "Ltd), the country's oldest and largest housing-finance company. At roughly USD 40 billion it was "
    "the largest M&A transaction in Indian corporate history. The merger became effective on "
    "1 July 2023, and HDFC Ltd shares were delisted shortly thereafter."
)

subhead("4.2 Deal Structure and Share-Swap Ratio")
make_table(
    ["Parameter", "Detail"],
    [
        ["Type", "Composite scheme of amalgamation (all-stock merger)"],
        ["Share-swap ratio", "42 HDFC Bank shares (face value Re 1) for every 25 HDFC Ltd shares "
         "(face value Rs 2) \u2014 a ratio of 1:1.68"],
        ["Announcement", "4 April 2022"],
        ["Effective date", "1 July 2023"],
        ["Shares allotted", "Over 311 crore new HDFC Bank shares"],
        ["Post-merger ownership", "HDFC Bank becomes 100% publicly owned; former HDFC Ltd shareholders "
         "hold ~41% of the bank"],
    ],
    widths=[1.8, 4.2],
)
caption("Sources: HDFC Bank press releases, Economic Times, Reuters.")

subhead("4.3 Strategic Rationale")
bullet("Create a larger balance sheet and net worth to support bigger-ticket lending, including "
       "infrastructure loans.")
bullet("Fund HDFC Ltd's mortgage book with HDFC Bank's low-cost CASA (current and savings account) "
       "deposit franchise, reducing the cost of funds.")
bullet("Cross-sell mortgages to the bank's large customer base and bank products to HDFC's mortgage "
       "customers.")
bullet("Simplify the group structure and remove the regulatory and operational friction of a "
       "bank lending through its parent.")

subhead("4.4 Synergies and Value Creation")
body(
    "Revenue synergies arise from cross-selling across a vastly larger combined customer base. Cost "
    "synergies come from cheaper deposit-based funding replacing wholesale borrowing, and from scale "
    "economies. Financial synergies stem from a stronger, more diversified balance sheet able to "
    "underwrite larger loans. The combination created the world's fourth-largest bank by market "
    "capitalisation (around USD 172 billion at completion), with assets of roughly USD 390 billion."
)

subhead("4.5 Market Reaction and Benefits to Shareholders")
body(
    "The announcement was received positively, with both stocks rising sharply on the news. The "
    "favourable swap ratio also created a well-publicised arbitrage opportunity for investors ahead "
    "of completion. HDFC Ltd shareholders received liquid, regulated banking-sector equity in "
    "exchange for housing-finance shares, while HDFC Bank shareholders gained a captive mortgage "
    "engine and a deeper, more diversified franchise."
)

subhead("4.6 Post-Merger Performance and Analysis")
body(
    "The integration enlarged the bank's balance sheet substantially and reshaped its funding and "
    "lending mix. The central post-merger challenge has been managing the cost of funds and the "
    "loan-to-deposit ratio inherited from HDFC Ltd's borrowing-funded mortgage book \u2014 making "
    "deposit mobilisation the key metric for judging whether the anticipated synergies are realised "
    "over time."
)

# ============================================================
# SECTION 5 - COMPARATIVE EXIT ANALYSIS
# ============================================================
heading("Section 5: Comparative Private Equity Exit Analysis", 1)
body(
    "The cases above illustrate different monetisation and consolidation routes. The table below "
    "compares the five principal PE exit strategies across the dimensions that matter most to a fund."
)
make_table(
    ["Exit Route", "Advantages", "Disadvantages", "Typical Timing / Return"],
    [
        ["IPO (incl. partial / OFS)",
         "High valuation in strong markets; retains upside; prestige",
         "Market-timing risk; lock-ins; phased exit",
         "Bull-market window; high return potential (e.g., Hexaware)"],
        ["Strategic sale (trade sale)",
         "Control premium; synergies; clean full exit",
         "Fewer buyers; antitrust review",
         "When a motivated strategic buyer exists; strong return"],
        ["Secondary buyout (sponsor-to-sponsor)",
         "Speed; certainty; confidentiality",
         "Buyer also needs return, capping price",
         "Mid-cycle; moderate return"],
        ["Partial exit",
         "Liquidity now + retained upside; de-risks",
         "Capital still tied up; residual market risk",
         "When further value is expected; balanced return"],
        ["Recapitalisation",
         "Returns cash without selling; keeps control",
         "Adds leverage and risk",
         "When credit is cheap; moderate return"],
    ],
    widths=[1.4, 1.9, 1.7, 1.6],
)

# ============================================================
# SECTION 6 - FINDINGS & RECOMMENDATIONS
# ============================================================
heading("Section 6: Key Findings and Recommendations", 1)

subhead("6.1 Lessons from Hexaware (Carlyle)")
bullet("A disciplined value-creation plan \u2014 growth investment, AI capability and margin expansion "
       "\u2014 roughly doubled revenue over the hold period.")
bullet("A partial IPO exit via offer-for-sale monetised gains while retaining majority upside, "
       "demonstrating flexible exit design.")
bullet("Listing into a strong IPO window maximised realised valuation.")

subhead("6.2 Lessons from the HDFC\u2013HDFC Bank Merger")
bullet("Synergy logic must be concrete \u2014 here, low-cost deposits funding a large mortgage book.")
bullet("An all-stock structure with a clearly justified swap ratio aligns both shareholder groups.")
bullet("Scale alone is not value; realising deposit-led funding synergies is the true test.")

subhead("6.3 Exit-Strategy Comparison \u2014 Key Takeaway")
body(
    "No single exit route is universally superior. The optimal route depends on market conditions, "
    "the availability of strategic buyers, the fund's need for liquidity versus continued upside, and "
    "the maturity of the asset. Hexaware shows that a partial IPO can be ideal for a high-growth, "
    "listable asset, whereas a trade sale better suits assets whose value is highest to a strategic "
    "acquirer."
)

subhead("6.4 Recommended Exit Strategy for Siemens India")
body(
    "For a Siemens India-type asset \u2014 a high-quality, strategically controlled industrial business "
    "\u2014 the most realistic and value-accretive route is continued value unlocking through demerger "
    "and listing of distinct businesses (as already executed for the energy unit), complemented by "
    "selective secondary stake sales to improve free float and, where appropriate, carve-outs of "
    "non-core units to financial buyers. A full PE-style buyout is improbable given the parent's "
    "strategic control."
)

subhead("6.5 Conclusion")
body(
    "Private equity value creation and M&A are two sides of the same coin: both seek to make a "
    "business worth more than it is today and to capture that value through a well-structured, well-"
    "timed transaction. The Hexaware case demonstrates disciplined value creation and a flexible "
    "partial exit; the HDFC merger demonstrates synergy-driven consolidation at national scale; and "
    "the Siemens India analysis shows how the same exit toolkit can be applied prospectively. "
    "Together they confirm that thesis clarity, operational improvement, sound deal structuring and "
    "exit timing are the enduring determinants of value creation."
)

# ============================================================
# REFERENCES
# ============================================================
heading("References and Data Sources", 1)
refs = [
    "Carlyle Group \u2014 press release on the listing of Hexaware Technologies (carlyle.com, Feb 2025).",
    "Reuters \u2014 reporting on the Carlyle\u2013Hexaware acquisition and the IPO offer-for-sale.",
    "Moneycontrol, CNBC-TV18, NDTV Profit, Livemint, Economic Times, Financial Express \u2014 coverage of "
    "the Hexaware acquisition, delisting and 2025 IPO.",
    "Hexaware Technologies \u2014 quarterly and annual results (CY2024, CY2025) via PRNewswire and company site.",
    "Stockopedia / company filings \u2014 Hexaware historical revenue data.",
    "HDFC Bank \u2014 press releases on the merger announcement (April 2022) and effectiveness (July 2023).",
    "Reuters, Economic Times, Onmanorama, NDTV Profit, The Economist \u2014 coverage of the HDFC\u2013HDFC Bank merger.",
    "Siemens Limited \u2014 press releases on quarterly results and the energy-business demerger (siemens.com).",
    "General references: standard finance texts on private equity, M&A and corporate valuation.",
]
for r in refs:
    bullet(r)

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(10)
run = note.add_run(
    "Disclaimer: This report is prepared for academic purposes. Financial figures are drawn from "
    "publicly available reporting and have been rounded; they should be verified against primary "
    "filings before any decision-making. The Siemens India section is an analytical / hypothetical "
    "exercise and does not assert any actual or planned transaction."
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = GREY

# ----- footer with page numbers -----
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
fldChar1 = OxmlElement("w:fldChar")
fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.set(qn("xml:space"), "preserve")
instrText.text = "PAGE"
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "end")
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

out = "/projects/sandbox/PE_Exit_and_MA_Case_Study_Report.docx"
doc.save(out)
print("Saved:", out)
print("Paragraphs:", len(doc.paragraphs), "| Tables:", len(doc.tables))
